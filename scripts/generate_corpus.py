"""
Generate the synthetic document corpus for Nilachala Textiles.

Reads:  data/corpus_manifest.csv  +  scripts/corpus_content.py
Writes: data/raw/*.pdf, *.docx

Run from the project root:
    python scripts/generate_corpus.py
"""

import csv
import io
import random
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)

from docx import Document
from docx.shared import Pt

from pdf2image import convert_from_path
from PIL import Image, ImageFilter

from corpus_content import CORPUS, FOOTER_NOTE

ROOT = Path(__file__).resolve().parent.parent
MANIFEST = ROOT / "data" / "corpus_manifest.csv"
OUTPUT_DIR = ROOT / "data" / "raw"
TMP_DIR = ROOT / "data" / ".tmp"

random.seed(42)

styles = getSampleStyleSheet()
styles.add(ParagraphStyle("DocTitle", parent=styles["Title"], fontSize=20, spaceAfter=6))
styles.add(ParagraphStyle("Meta", parent=styles["Normal"], fontSize=8,
                          textColor=colors.grey, spaceAfter=2))
styles.add(ParagraphStyle("Caption", parent=styles["Normal"], fontSize=8,
                          fontName="Helvetica-Oblique", spaceAfter=3))
styles.add(ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10,
                          leading=15, spaceAfter=6, alignment=4))


# ----------------------------------------------------------------------
# helpers
# ----------------------------------------------------------------------

def load_manifest():
    with open(MANIFEST, newline="", encoding="utf-8") as f:
        return {r["doc_id"]: r for r in csv.DictReader(f)}


def is_true(value):
    """csv.DictReader gives strings; 'FALSE' is truthy, so compare explicitly."""
    return str(value).strip().upper() == "TRUE"


def page_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.grey)
    canvas.drawString(20 * mm, 12 * mm, FOOTER_NOTE)
    canvas.drawRightString(A4[0] - 20 * mm, 12 * mm, f"Page {canvas.getPageNumber()}")
    canvas.restoreState()


def make_table(spec):
    header = [Paragraph(f"<b>{c}</b>", styles["Meta"]) for c in spec["columns"]]
    body = [[Paragraph(str(c), styles["Meta"]) for c in row] for row in spec["rows"]]
    tbl = Table([header] + body, repeatRows=1, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e4e4e4")),
        ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#666666")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f6f6f6")]),
    ]))
    return tbl


def build_flow(doc_id, row, target_pages):
    """Build the reportlab flowable list for one document."""
    content = CORPUS[doc_id]
    flow = [
        Paragraph(content["title"], styles["DocTitle"]),
        Paragraph("Nilachala Textiles Pvt. Ltd., Bhubaneswar", styles["Meta"]),
        Paragraph(
            f"Document ref: {doc_id} &nbsp;&nbsp;|&nbsp;&nbsp; "
            f"Version {row['version']} &nbsp;&nbsp;|&nbsp;&nbsp; "
            f"Owner: {row['department']} department",
            styles["Meta"],
        ),
        Spacer(1, 8 * mm),
    ]

    for section in content["sections"]:
        flow.append(Paragraph(section["heading"], styles["Heading2"]))
        if "table" in section:
            spec = section["table"]
            flow.append(Paragraph(spec["caption"], styles["Caption"]))
            flow.append(make_table(spec))
            flow.append(Spacer(1, 5 * mm))
        for para in section.get("body", []):
            flow.append(Paragraph(para, styles["Body"]))
        flow.append(Spacer(1, 3 * mm))

    # Employee Handbook needs to be long enough to stress the chunker.
    if target_pages >= 60:
        flow.extend(_appendix_flow())

    return flow


def _appendix_flow():
    """Filler appendices so D03 reaches realistic handbook length."""
    flow = [PageBreak(), Paragraph("Appendix A - Departmental Standing Instructions",
                                   styles["Heading1"])]
    departments = [
        "Spinning", "Weaving", "Dyeing and Finishing", "Quality Control",
        "Packing and Despatch", "Stores and Inventory", "Maintenance",
        "Human Resources", "Finance and Accounts", "Information Technology",
        "Security", "Canteen and Welfare",
    ]
    topics = [
        ("Reporting lines",
         "Each employee reports to a single designated supervisor. Where an employee "
         "works across shifts, the shift supervisor on duty holds day to day authority "
         "while the functional head retains responsibility for performance assessment."),
        ("Shift handover",
         "A written handover must be completed at the end of every shift covering work "
         "completed, work in progress, outstanding issues and any safety observation. "
         "The incoming supervisor countersigns the handover record."),
        ("Material requisition",
         "Materials are drawn from stores against an approved requisition. Requisitions "
         "above the departmental limit require the countersignature of the department "
         "head. Emergency issues are permitted against a verbal approval which must be "
         "regularised in writing within one working day."),
        ("Quality escalation",
         "Any deviation from specification must be recorded and escalated to the Quality "
         "Control department within the same shift. Production may not continue on a "
         "batch under quality hold without written clearance."),
        ("Record retention",
         "Departmental records are retained for a minimum of three years. Statutory "
         "records are retained for the period prescribed by the relevant legislation. "
         "Disposal of records requires the approval of the department head."),
    ]
    for dept in departments:
        flow.append(Paragraph(f"A.{departments.index(dept) + 1} {dept} Department",
                              styles["Heading2"]))
        for heading, text in topics:
            flow.append(Paragraph(heading, styles["Heading3"]))
            flow.append(Paragraph(
                text.replace("Each employee", f"Each {dept} employee", 1),
                styles["Body"]))
        flow.append(Spacer(1, 4 * mm))
    return flow


# ----------------------------------------------------------------------
# renderers
# ----------------------------------------------------------------------

def render_pdf(doc_id, row, out_path, target_pages):
    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=20 * mm, bottomMargin=22 * mm,
        title=CORPUS[doc_id]["title"], author="Nilachala Textiles Pvt. Ltd.",
    )
    doc.build(build_flow(doc_id, row, target_pages),
              onFirstPage=page_footer, onLaterPages=page_footer)


def render_docx(doc_id, row, out_path):
    content = CORPUS[doc_id]
    d = Document()

    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    d.add_heading(content["title"], level=0)
    meta = d.add_paragraph()
    meta.add_run("Nilachala Textiles Pvt. Ltd., Bhubaneswar\n").italic = True
    meta.add_run(
        f"Document ref: {doc_id}  |  Version {row['version']}  |  "
        f"Owner: {row['department']} department"
    ).italic = True

    for section in content["sections"]:
        d.add_heading(section["heading"], level=1)
        if "table" in section:
            spec = section["table"]
            cap = d.add_paragraph(spec["caption"])
            cap.runs[0].italic = True
            t = d.add_table(rows=1, cols=len(spec["columns"]))
            t.style = "Table Grid"
            for i, col in enumerate(spec["columns"]):
                cell = t.rows[0].cells[i]
                cell.text = col
                cell.paragraphs[0].runs[0].bold = True
            for row_vals in spec["rows"]:
                cells = t.add_row().cells
                for i, val in enumerate(row_vals):
                    cells[i].text = str(val)
            d.add_paragraph()
        for para in section.get("body", []):
            d.add_paragraph(para)

    d.add_paragraph()
    footer = d.add_paragraph(FOOTER_NOTE)
    footer.runs[0].italic = True
    d.save(str(out_path))


def render_scanned_pdf(doc_id, row, out_path, target_pages):
    """
    Render normally, rasterise to images, degrade, rebuild as an image-only PDF.
    Result has no extractable text layer - forces the OCR path.
    """
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    tmp_pdf = TMP_DIR / f"{doc_id}_clean.pdf"
    render_pdf(doc_id, row, tmp_pdf, target_pages)

    pages = convert_from_path(str(tmp_pdf), dpi=150)

    degraded = []
    for i, page in enumerate(pages):
        img = page.convert("L")                                  # greyscale
        img = img.rotate(random.uniform(-0.7, 0.7),              # slight skew
                         expand=True, fillcolor=255)
        img = img.filter(ImageFilter.GaussianBlur(radius=0.4))   # soft focus

        # JPEG round trip introduces scan-like compression artefacts
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=72)
        buf.seek(0)
        img = Image.open(buf).convert("RGB")

        degraded.append(img)

    degraded[0].save(str(out_path), save_all=True, append_images=degraded[1:],
                     resolution=150.0)
    tmp_pdf.unlink(missing_ok=True)


# ----------------------------------------------------------------------
# main
# ----------------------------------------------------------------------

def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    manifest = load_manifest()

    for doc_id, row in manifest.items():
        if doc_id not in CORPUS:
            print(f"[skip] {doc_id} present in manifest but missing from CORPUS")
            continue

        out_path = OUTPUT_DIR / row["filename"]
        target_pages = int(row["pages"])
        fmt = row["format"].strip().lower()

        if fmt == "docx":
            render_docx(doc_id, row, out_path)
            kind = "docx"
        elif is_true(row["is_scanned"]):
            render_scanned_pdf(doc_id, row, out_path, target_pages)
            kind = "pdf (scanned)"
        else:
            render_pdf(doc_id, row, out_path, target_pages)
            kind = "pdf"

        size_kb = out_path.stat().st_size / 1024
        print(f"[ok] {doc_id}  {kind:14s} {size_kb:8.1f} KB  {row['filename']}")

    if TMP_DIR.exists():
        for leftover in TMP_DIR.glob("*"):
            leftover.unlink()
        TMP_DIR.rmdir()

    print(f"\nDone. Files written to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
